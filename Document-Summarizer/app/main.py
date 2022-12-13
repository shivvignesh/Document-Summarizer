from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from app.nlp import NLP
from pprint import pprint
from docx import *
from flask import *

app1 = Flask(__name__)  

class Message(BaseModel):
    input: str = None
    output: str = None
    sections: list = None

app = FastAPI()
nlp = NLP()

origins = [
    "http://localhost",
    "http://localhost:3000",
    "http://127.0.0.1:3000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["POST"],
    allow_headers=["*"],
)

@app.post("/generative/")
async def  generate(message: Message):
    message.output  = nlp.generate(prompt=message.input)
    return {"output" : message.output}

@app.post("/sentiment/")
async def sentiment_analysis(message: Message):
    message.output  = str(nlp.sentiments(message.input))
    return {"output" : message.output}

@app.post('/add_document/')
async def read_doc(file: UploadFile = File(...)):

    # filename = file.filename
    # bolds = []
    # document = Document(file)

    # for para in document.paragraphs:
    #     for run in para.runs:
    #         if run.bold:
    #             bolds.append(run.text)

    # print("Bold responses "+str(bolds))

    return {"outputs" : file.filename}


    # if request.method == 'POST':
    #     file = request.files['file']
        # filename = file.filename
        # bolds = []
        # document = Document(file)

        # for para in document.paragraphs:
        #     for run in para.runs:
        #         if run.bold:
        #             bolds.append(run.text)

        # print("Bold responses "+str(bolds))

        # return {"outputs" : "success"}


@app1.route('/')  
def upload(message: Message):  
    return render_template("file_upload_form.html") 


if __name__ == '__main__':  
    app.run(debug = True)  