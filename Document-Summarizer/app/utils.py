from docx import *
from flask import *
import os

app = Flask(__name__)

ALLOWED_EXTENSIONS = {'docx'}

uploads_dir = os.path.join(app.instance_path, 'uploads')
os.makedirs(uploads_dir, exist_ok=True)

def get_content(document):
    doc_name = document+ '.docx'
    document = Document(os.path.join(uploads_dir,doc_name))
    bolds = []
    for para in document.paragraphs:
            for run in para.runs:
                if run.bold:
                    bolds.append(run.text)
    
    mytext = ""
    for para in document.paragraphs:
        mytext = mytext + para.text
        
    sections = {}

    for i in range(len(bolds)-1):
        content = mytext[mytext.index(bolds[i]):mytext.index(bolds[i+1])]
        content = content.replace(bolds[i],"")
        sections[bolds[i]] = content

    return sections,bolds




def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS